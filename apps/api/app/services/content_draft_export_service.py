"""Render content drafts to downloadable .txt / .docx files.

Lets an operator pull a whole social pack (or any single draft) out as a Word
document or plain-text file instead of copying each post by hand. The .docx
embeds the generated creative image when present; .txt links to it.

Mirrors `suggested_copy_service`'s rendering conventions and reuses its
filename + media-type helpers. Images are fetched through the SSRF guard
because `content_drafts.image_url` is an operator-editable field.
"""

from __future__ import annotations

from io import BytesIO
from uuid import UUID

from sqlalchemy.orm import Session

from app.core.logging import get_logger
from app.models.content_draft import ContentDraft
from app.services.suggested_copy_service import DOCX_MEDIA_TYPE, safe_filename

log = get_logger(__name__)

__all__ = [
    "DOCX_MEDIA_TYPE",
    "safe_filename",
    "get_drafts_by_ids",
    "render_txt",
    "render_bundle_txt",
    "render_docx",
    "render_bundle_docx",
]

# Cap on an image we'll pull into a .docx, so a hostile/huge image_url can't
# balloon the document or memory.
_MAX_IMAGE_BYTES = 10 * 1024 * 1024


def get_drafts_by_ids(
    db: Session, *, workspace_id: UUID, ids: list[UUID]
) -> list[ContentDraft]:
    """Workspace-scoped fetch preserving the caller's id order (so a downloaded
    pack keeps the on-screen order)."""
    if not ids:
        return []
    rows = (
        db.query(ContentDraft)
        .filter(
            ContentDraft.workspace_id == workspace_id,
            ContentDraft.id.in_(ids),
        )
        .all()
    )
    by_id = {r.id: r for r in rows}
    return [by_id[i] for i in ids if i in by_id]


# ---------------------------------------------------------------------------
# Shared label helpers
# ---------------------------------------------------------------------------


def _type_label(draft: ContentDraft) -> str:
    return "Short video script" if draft.type.value == "short_video_script" else "Post"


def _meta_line(draft: ContentDraft) -> str:
    bits = [_type_label(draft)]
    if draft.platform:
        label = (draft.seo_metadata or {}).get("platform_label") or draft.platform
        bits.append(str(label))
    return "  ·  ".join(bits)


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


def _txt_block(draft: ContentDraft) -> str:
    lines = [
        draft.title,
        "=" * len(draft.title),
        _meta_line(draft),
        "",
        (draft.body or "").strip(),
    ]
    if draft.hashtags:
        lines += ["", " ".join(draft.hashtags)]
    if draft.keywords:
        lines += ["", f"Keywords: {', '.join(draft.keywords)}"]
    if draft.image_url:
        lines += ["", f"Image: {draft.image_url}"]
    return "\n".join(lines) + "\n"


def render_txt(draft: ContentDraft) -> bytes:
    return _txt_block(draft).encode("utf-8")


def render_bundle_txt(drafts: list[ContentDraft], *, title: str) -> bytes:
    header = f"{title}\n{'#' * 60}\n{len(drafts)} item(s)\n\n"
    body = ("\n" + "-" * 60 + "\n\n").join(_txt_block(d) for d in drafts)
    return (header + body).encode("utf-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _new_document():
    from docx import Document  # lazy import so a missing dep only hits DOCX

    return Document()


def _write_body_to_doc(doc, body: str) -> None:
    """Light-markdown body → Word paragraphs (## headings, - bullets)."""
    for raw_line in (body or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)


def _fetch_image_bytes(draft: ContentDraft) -> bytes | None:
    """Return the draft's image bytes, or None if absent/unavailable.

    Local `/uploads/...` paths are read from disk (guarded against traversal);
    absolute URLs are fetched through the SSRF guard, since `image_url` is an
    operator-editable field and could otherwise point at an internal host."""

    url = (draft.image_url or "").strip()
    if not url:
        return None

    if url.startswith("/uploads/"):
        from app.services.image_upload_service import uploads_root

        root = uploads_root().resolve()
        candidate = (root / url[len("/uploads/") :]).resolve()
        if not candidate.is_relative_to(root) or not candidate.is_file():
            return None
        data = candidate.read_bytes()
        return data[:_MAX_IMAGE_BYTES] if data else None

    if url.startswith(("http://", "https://")):
        from app.security.safe_http import safe_get

        try:
            resp = safe_get(url, timeout=15.0)
        except Exception as exc:  # noqa: BLE001 — blocked/unreachable → skip image
            log.info("content_export.image_fetch_failed", url=url, error=str(exc))
            return None
        if resp.status_code >= 400:
            return None
        return resp.content[:_MAX_IMAGE_BYTES] or None

    return None


def _add_draft_to_doc(doc, draft: ContentDraft) -> None:
    from docx.shared import Inches

    doc.add_heading(draft.title, level=1)
    doc.add_paragraph(_meta_line(draft)).italic = True

    image_bytes = _fetch_image_bytes(draft)
    if image_bytes:
        try:
            doc.add_picture(BytesIO(image_bytes), width=Inches(5.5))
        except Exception as exc:  # noqa: BLE001 — unrecognized image → skip, keep text
            log.info("content_export.image_embed_failed", error=str(exc))

    _write_body_to_doc(doc, draft.body or "")

    if draft.hashtags:
        doc.add_paragraph("")
        doc.add_paragraph(" ".join(draft.hashtags))
    if draft.keywords:
        p = doc.add_paragraph()
        p.add_run(f"Keywords: {', '.join(draft.keywords)}").italic = True


def render_docx(draft: ContentDraft) -> bytes:
    doc = _new_document()
    _add_draft_to_doc(doc, draft)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_bundle_docx(drafts: list[ContentDraft], *, title: str) -> bytes:
    doc = _new_document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"{len(drafts)} item(s)")
    for idx, draft in enumerate(drafts):
        if idx:
            doc.add_page_break()
        _add_draft_to_doc(doc, draft)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
